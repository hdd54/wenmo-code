"""Word 文档生成插件：用 python-docx 正确生成 .docx（中文编码可靠，支持标题/段落/列表/表格）。
模型不要手写 .docx（二进制格式会损坏/乱码）。"""

import os
import re
import time

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FILES_DIR = os.path.join(BASE, "files")
os.makedirs(FILES_DIR, exist_ok=True)
MML2OMML_XSL = os.path.join(BASE, "MML2OMML.XSL")


def _latex_to_omml(latex):
    """LaTeX 公式 → MathML（latex2mathml）→ OMML（MML2OMML.XSL）→ Word 原生公式（MathType 兼容）。失败返回 None。"""
    try:
        import latex2mathml.converter
        import lxml.etree as ET
        mathml = latex2mathml.converter.convert(latex)
        if not os.path.isfile(MML2OMML_XSL):
            return None
        xsl = ET.parse(MML2OMML_XSL)
        transform = ET.XSLT(xsl)
        mml_tree = ET.fromstring(mathml.encode("utf-8"))
        omml_tree = transform(mml_tree)
        s = str(omml_tree)
        return s if "<m:oMath" in s else None
    except Exception:
        return None


def _insert_omml(paragraph, omml_str):
    """把 OMML XML 追加到段落（Word 原生公式）"""
    import lxml.etree as ET
    paragraph._p.append(ET.fromstring(omml_str.encode("utf-8")))


def _apply_doc_fonts(doc):
    """文档默认字体：中文宋体、英文 Times New Roman（正文 12pt；标题黑体）"""
    try:
        from docx.oxml.ns import qn
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = docx.shared.Pt(12)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "宋体")
        for hn in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            try:
                hs = doc.styles[hn]
                hs.font.name = "Times New Roman"
                hrpr = hs.element.get_or_add_rPr()
                hrfonts = hrpr.find(qn("w:rFonts"))
                if hrfonts is None:
                    hrfonts = hrpr.makeelement(qn("w:rFonts"), {})
                    hrpr.append(hrfonts)
                hrfonts.set(qn("w:eastAsia"), "黑体")
            except Exception:
                continue
    except Exception:
        pass


def _safe(s, n=60):
    s = re.sub(r"[\\/:*?\"<>|]", "_", str(s or "文档"))
    return s[:n] or "文档"


def _ensure_run_fonts(doc):
    """run 级字体强制（样式继承不稳）：正文中文宋体 + 英文 Times New Roman；标题中文黑体"""
    try:
        from docx.oxml.ns import qn
        for para in doc.paragraphs:
            is_heading = bool(para.style and para.style.name and "Heading" in para.style.name)
            cn = "黑体" if is_heading else "宋体"
            for run in para.runs:
                if run._element.find(qn("w:rPr")) is not None and run._element.rPr.find(qn("w:rFonts")) is not None:
                    continue   # 已有显式字体（如公式 run）不动
                run.font.name = "Times New Roman"
                rpr = run._element.get_or_add_rPr()
                rf = rpr.find(qn("w:rFonts"))
                if rf is None:
                    rf = rpr.makeelement(qn("w:rFonts"), {})
                    rpr.append(rf)
                rf.set(qn("w:eastAsia"), cn)
                rf.set(qn("w:ascii"), "Times New Roman")
                rf.set(qn("w:hAnsi"), "Times New Roman")
    except Exception:
        pass


def _resolve_image_local(src):
    import base64
    import tempfile
    import urllib.request
    s = str(src or "").strip()
    if not s:
        return None
    if s.startswith("data:image/"):
        m = re.match(r"^data:image/(\w+);base64,(.+)$", s, re.S)
        if not m:
            return None
        ext = {"png": "png", "jpeg": "jpg", "jpg": "jpg", "gif": "gif", "webp": "webp"}.get(m.group(1).lower(), "png")
        try:
            raw = base64.b64decode(m.group(2))
            fd, tmp = tempfile.mkstemp(suffix="." + ext)
            with os.fdopen(fd, "wb") as f:
                f.write(raw)
            return tmp
        except Exception:
            return None
    if s.startswith("http://") or s.startswith("https://"):
        try:
            req = urllib.request.Request(s, headers={"User-Agent": "Mozilla/5.0"})
            raw = urllib.request.urlopen(req, timeout=15).read()
            fd, tmp = tempfile.mkstemp(suffix=".jpg")
            with os.fdopen(fd, "wb") as f:
                f.write(raw)
            return tmp
        except Exception:
            return None
    p = s
    if p.startswith("/files/"):
        p = os.path.join(FILES_DIR, os.path.basename(p))
    return p if os.path.isfile(p) else None


def create_docx(args):
    """用 python-docx 生成 Word 文档（标题/段落/列表/表格），保存到下载区并返回链接。"""
    title = str(args.get("title", "")).strip() or "文档"
    subtitle = str(args.get("subtitle", "")).strip() or ""
    blocks = args.get("blocks") or args.get("content") or []
    if not isinstance(blocks, list) or not blocks:
        return {"error": "blocks 参数不能为空：请提供 [{type: heading/paragraph/list/table, ...}] 列表"}
    try:
        import docx
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"error": "服务器缺少 python-docx，无法生成 Word"}
    try:
        doc = docx.Document()
        _apply_doc_fonts(doc)   # 中文宋体 + 英文 Times New Roman（标题黑体）
        # 标题
        h = doc.add_heading(title, level=0)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            r.font.size = Pt(26)
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.size = Pt(13)
        for b in blocks:
            if not isinstance(b, dict):
                continue
            btype = str(b.get("type", "paragraph")).lower()
            text = str(b.get("text", "")).strip()
            if btype == "heading":
                level = min(4, max(1, int(b.get("level", 1) or 1)))
                doc.add_heading(text or "", level=level)
            elif btype in ("h1", "h2", "h3"):
                doc.add_heading(text or "", level={"h1": 1, "h2": 2, "h3": 3}[btype])
            elif btype == "list":
                items = b.get("items") or ([text] if text else [])
                for it in items:
                    doc.add_paragraph(str(it), style="List Bullet")
            elif btype == "number":
                items = b.get("items") or ([text] if text else [])
                for it in items:
                    doc.add_paragraph(str(it), style="List Number")
            elif btype == "table":
                rows = b.get("rows") or []
                if rows and isinstance(rows, list):
                    ncols = max(len(r) for r in rows if isinstance(r, list))
                    table = doc.add_table(rows=len(rows), cols=max(1, ncols))
                    table.style = "Light Grid Accent 1"
                    for ri, row in enumerate(rows):
                        if not isinstance(row, list):
                            continue
                        for ci, cell in enumerate(row[:ncols]):
                            table.cell(ri, ci).text = str(cell or "")
            elif btype == "quote":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif btype in ("formula", "math"):
                # 公式块：LaTeX → OMML（Word 原生公式，MathType 兼容）；转换失败回退居中文本
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    omml = _latex_to_omml(text)
                    if omml:
                        _insert_omml(p, omml)
                    else:
                        run = p.add_run(text or "")
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(13)
                        run.italic = True
                except Exception:
                    run = p.add_run(text or "")
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(13)
                    run.italic = True
            elif btype == "image":
                # 图片块：src 支持 http(s) URL / data URL(base64) / 本地路径 / /files/ 路径
                src = str(b.get("src") or b.get("image") or b.get("url") or "").strip()
                if src:
                    from docx.shared import Inches
                    img_path = _resolve_image_local(src)
                    if img_path:
                        try:
                            w = float(b.get("width") or 5.5)
                            doc.add_picture(img_path, width=Inches(min(w, 6.5)))
                            cap = str(b.get("caption") or "").strip()
                            if cap:
                                cp = doc.add_paragraph()
                                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                r = cp.add_run(cap)
                                r.font.size = Pt(10)
                                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        except Exception as pic_err:
                            doc.add_paragraph("[图片插入失败: %s]" % pic_err)
                        finally:
                            # 清理下载的临时文件（本地路径不动）
                            if img_path.startswith(tempfile.gettempdir()):
                                try:
                                    os.remove(img_path)
                                except Exception:
                                    pass
                    else:
                        doc.add_paragraph("[图片无效：无法从该地址获取图片]")
            else:  # paragraph
                doc.add_paragraph(text)
        fname = f"{_safe(title)}_{int(time.time())}.docx"
        fpath = os.path.join(FILES_DIR, fname)
        _ensure_run_fonts(doc)   # run 级字体强制（宋体/Times）
        doc.save(fpath)
        url = f"http://127.0.0.1:8000/files/{fname}"
        return {
            "ok": True,
            "file": fname,
            "url": url,
            "note": f"Word 文档已生成：{fname}。请在回答里引用链接，用户点击即可下载（用 Word/WPS 打开排版正常）。",
        }
    except Exception as e:
        return {"error": f"Word 生成失败: {e}"}


PLUGIN_TOOLS = [
    {
        "name": "create_docx",
        "description": "创建 Word 文档（服务器 python-docx 生成，中文编码可靠；支持标题/段落/列表/编号/表格/引用，排版规范）。"
                       "当用户要求『做一个 Word 文档/报告/方案/合同/说明书』时，必须使用本工具，"
                       "不要自己手写 .docx（二进制格式会损坏/乱码）。"
                       "参数：title=文档标题；subtitle=副标题（可选）；"
                       "blocks=[{type: heading/paragraph/list/number/table/quote, ...}]。"
                       "生成后返回链接，请在回答中引用。",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "文档标题"},
                "subtitle": {"type": "string", "description": "副标题（居中）"},
                "blocks": {
                    "type": "array",
                    "items": {"type": "object"},
                        "description": "内容块列表：[{type:'heading',text,level} | {type:'paragraph',text} | "
                                       "{type:'list'/'number',items:[...]} | {type:'table',rows:[[...]]} | "
                                       "{type:'quote',text} | {type:'formula',text:LaTeX/数学符号} | "
                                       "{type:'image',src:图片URL/base64/路径,caption:图注,width:宽度英寸}]。"
                                       "公式用 formula 块（LaTeX 或 Unicode 数学符号，居中显示）。"
                }
            },
            "required": ["title", "blocks"],
        },
        "handler": create_docx,
    }
]
