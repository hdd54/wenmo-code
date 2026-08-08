"""生成一份排版好的个人简历 Word 文档 (.docx)。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn




def _set_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element.rPr
    if r is not None:
        rFonts = r.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), name)


def _heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=14, bold=True, color=(0x1F, 0x4E, 0x79))
    p.space_after = Pt(4)
    # 底部边框线
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pPr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '1', qn('w:color'): '1F4E79'})
    pBdr.append(bottom)
    pPr.append(pBdr)


def _item(doc, text, size=10.5, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def _kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    _set_font(r1, size=10.5, bold=True)
    r2 = p.add_run("　" + value)
    _set_font(r2, size=10.5)
    return p


def generate_resume_docx(path="个人简历.docx"):
    doc = Document()
    # 页边距
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("个 人 简 历")
    _set_font(run, size=22, bold=True, color=(0x1F, 0x4E, 0x79))

    # 基本信息
    _heading(doc, "基本信息")
    _kv(doc, "姓名：", "[你的姓名]")
    _kv(doc, "求职意向：", "[目标岗位]")
    _kv(doc, "电话：", "[手机号码]")
    _kv(doc, "邮箱：", "[邮箱地址]")
    _kv(doc, "所在地：", "[城市]")

    # 教育背景
    _heading(doc, "教育背景")
    _item(doc, "[学校名称] · [学历]　　[入学时间] - [毕业时间]", bold=True)
    _item(doc, "[专业名称]", indent=0.5)
    _item(doc, "· [相关课程/方向，如：主修××，研究方向××]", indent=0.5)
    _item(doc, "· [在校荣誉，如：奖学金、优秀毕业生等]", indent=0.5)

    # 工作经历
    _heading(doc, "工作经历")
    _item(doc, "[公司名称] · [职位]　　[起始时间] - [结束时间]", bold=True)
    _item(doc, "· [职责1：用一句话描述核心工作，尽量量化结果]", indent=0.5)
    _item(doc, "· [职责2：说明解决的问题或达成的目标]", indent=0.5)
    _item(doc, "· [职责3：体现产出，如“提升效率 30%”“服务 5000+ 用户”]", indent=0.5)
    _item(doc, "[上一家公司名称] · [职位]　　[起始时间] - [结束时间]", bold=True)
    _item(doc, "· [职责1]", indent=0.5)
    _item(doc, "· [职责2]", indent=0.5)

    # 项目经验
    _heading(doc, "项目经验")
    _item(doc, "[项目名称]　　[起止时间]", bold=True)
    _item(doc, "· [项目背景：一句话说明要解决的问题]", indent=0.5)
    _item(doc, "· [我的角色：负责哪部分工作]", indent=0.5)
    _item(doc, "· [技术/成果：方案与量化效果]", indent=0.5)
    _item(doc, "[项目名称]　　[起止时间]", bold=True)
    _item(doc, "· [项目背景]", indent=0.5)
    _item(doc, "· [我的角色]", indent=0.5)
    _item(doc, "· [技术/成果]", indent=0.5)

    # 技能清单
    _heading(doc, "技能清单")
    _item(doc, "· 语言/框架：[如 Python、JavaScript、React、Vue 等]", indent=0.5)
    _item(doc, "· 工具/平台：[如 Git、Docker、Linux、MySQL 等]", indent=0.5)
    _item(doc, "· 其他：[如英语水平 CET-6、可流利读写技术文档等]", indent=0.5)

    # 自我评价
    _heading(doc, "自我评价 / 亮点")
    _item(doc, "· [亮点1：最突出的优势，如“能独立从 0 到 1 搭建系统”]", indent=0.5)
    _item(doc, "· [亮点2：性格或协作特质，如“抗压能力强、善于跨团队沟通”]", indent=0.5)
    _item(doc, "· [亮点3：职业态度，如“自驱力强、注重代码质量与可维护性”]", indent=0.5)

    # 保存
    out = path
    doc.save(out)
    return "已生成简历文档：" + os.path.abspath(out)

PLUGIN_TOOLS = [{
    "name": "generate_resume_docx",
    "description": "生成一份专业排版的个人简历 Word 文档，保存到工作区并返回路径。",
    "parameters": {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "输出文件名，默认 个人简历.docx"}
        }
    },
    "handler": generate_resume_docx,
}]
